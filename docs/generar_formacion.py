"""
Script para generar el documento de formacion Directus para Vivaz Clay Targets.
Ejecutar con: python generar_formacion.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUTPUT_PATH = r"C:\Users\Piensaenweb\Documents\Claude\platos-vivaz\docs\formacion-directus-vivaz.docx"

# ── Colores corporativos ────────────────────────────────────────────────────
COLOR_VERDE      = RGBColor(0x2E, 0x7D, 0x32)   # verde oscuro Vivaz
COLOR_VERDE_CLARO = RGBColor(0x81, 0xC7, 0x84)  # verde claro
COLOR_GRIS_TEXTO = RGBColor(0x33, 0x33, 0x33)
COLOR_GRIS_CLARO = RGBColor(0xF5, 0xF5, 0xF5)
COLOR_ADVERTENCIA = RGBColor(0xE5, 0x35, 0x35)  # rojo advertencia
COLOR_NOTA       = RGBColor(0xE6, 0x5F, 0x00)   # naranja nota

# ── Helpers de formato ─────────────────────────────────────────────────────

def set_cell_background(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_horizontal_line(doc, color_hex="2E7D32"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(0)
    return p

def heading_1(doc, text):
    """Titulo de seccion numerada principal."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = COLOR_VERDE
    # linea separadora bajo el titulo
    add_horizontal_line(doc, "2E7D32")
    return p

def heading_2(doc, text):
    """Sub-titulo de paso."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = COLOR_VERDE
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_GRIS_TEXTO
    return p

def step_list(doc, steps):
    """Lista numerada de pasos."""
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(step)
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_GRIS_TEXTO

def bullet_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_GRIS_TEXTO

def captura_placeholder(doc, descripcion):
    """Bloque visual que indica donde va una captura de pantalla."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_background(cell, "EEF7EE")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run(f"[CAPTURA: {descripcion}]")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_VERDE
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def nota(doc, texto, tipo="NOTA"):
    """Bloque de nota informativa o advertencia."""
    color = COLOR_ADVERTENCIA if tipo == "ADVERTENCIA" else COLOR_NOTA
    color_bg = "FFF3E0" if tipo == "NOTA" else "FFEBEE"
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_background(cell, color_bg)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.3)
    run_label = p.add_run(f"{tipo}: ")
    run_label.bold = True
    run_label.font.size = Pt(11)
    run_label.font.color.rgb = color
    run_text = p.add_run(texto)
    run_text.font.size = Pt(11)
    run_text.font.color.rgb = COLOR_GRIS_TEXTO
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def campo_tabla(doc, filas):
    """Tabla de dos columnas: Campo | Descripcion."""
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    # Cabecera
    hdr = table.rows[0].cells
    set_cell_background(hdr[0], "2E7D32")
    set_cell_background(hdr[1], "2E7D32")
    for cell, texto in zip(hdr, ["Campo", "Qué significa / qué escribir"]):
        p = cell.paragraphs[0]
        run = p.add_run(texto)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Filas
    for i, (campo, desc) in enumerate(filas):
        row = table.add_row().cells
        bg = "F5F5F5" if i % 2 == 0 else "FFFFFF"
        set_cell_background(row[0], bg)
        set_cell_background(row[1], bg)
        r0 = row[0].paragraphs[0].add_run(campo)
        r0.bold = True
        r0.font.size = Pt(10)
        r0.font.color.rgb = COLOR_VERDE
        r1 = row[1].paragraphs[0].add_run(desc)
        r1.font.size = Pt(10)
        r1.font.color.rgb = COLOR_GRIS_TEXTO
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ── Portada ────────────────────────────────────────────────────────────────

def add_portada(doc):
    # Espacio superior
    for _ in range(4):
        doc.add_paragraph()

    # Logo placeholder
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_background(cell, "2E7D32")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(20)
    run = p.add_run("[LOGO PIENSAENWEB]")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()

    # Titulo principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Guía de uso del gestor de contenidos")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = COLOR_VERDE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Vivaz Clay Targets")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = COLOR_GRIS_TEXTO

    doc.add_paragraph()
    add_horizontal_line(doc, "81C784")
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Directus CMS · Versión 1.0")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Mayo 2026  ·  Preparado por Piensaenweb para Vivaz")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.add_page_break()

# ── Indice ─────────────────────────────────────────────────────────────────

def add_indice(doc):
    heading_1(doc, "Índice de contenidos")
    secciones = [
        ("1.", "Cómo acceder al panel de administración"),
        ("2.", "Cómo editar un producto existente"),
        ("3.", "Cómo publicar una noticia nueva"),
        ("4.", "Cómo gestionar las traducciones"),
        ("5.", "Cómo ver los contactos recibidos (leads)"),
        ("6.", "Qué NO debes tocar"),
        ("7.", "Preguntas frecuentes y contacto de soporte"),
    ]
    for num, titulo in secciones:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run_num = p.add_run(f"{num}  ")
        run_num.bold = True
        run_num.font.size = Pt(12)
        run_num.font.color.rgb = COLOR_VERDE
        run_txt = p.add_run(titulo)
        run_txt.font.size = Pt(12)
        run_txt.font.color.rgb = COLOR_GRIS_TEXTO
    doc.add_page_break()

# ── Sección 1: Acceso ──────────────────────────────────────────────────────

def add_seccion_acceso(doc):
    heading_1(doc, "1. Cómo acceder al panel de administración")

    body(doc, (
        "El panel de administración es la herramienta desde la que gestionas todo el "
        "contenido de tu web: productos, noticias y contactos recibidos."
    ))

    heading_2(doc, "Datos de acceso")

    campo_tabla(doc, [
        ("URL del panel",      "[PLACEHOLDER: https://cms.vivazclat.com/admin]"),
        ("Usuario (email)",    "[PLACEHOLDER: juanma@vivazclat.com]"),
        ("Contraseña",         "[PLACEHOLDER: facilitada por Piensaenweb]"),
    ])

    nota(doc, (
        "Guarda estos datos en un lugar seguro. Si olvidas la contraseña, "
        "contacta con Piensaenweb para restablecerla."
    ))

    heading_2(doc, "Pasos para entrar")

    step_list(doc, [
        "Abre tu navegador (Chrome o Firefox recomendados).",
        "Escribe la URL del panel en la barra de direcciones y pulsa Enter.",
        "Introduce tu email y contraseña en el formulario de inicio de sesión.",
        "Haz clic en el botón «Iniciar sesión».",
        "Verás el panel principal (Dashboard) con el menú de colecciones a la izquierda.",
    ])

    captura_placeholder(doc, "Pantalla de inicio de sesión de Directus")
    captura_placeholder(doc, "Panel principal con el menú lateral visible")

    nota(doc, (
        "Si el panel tarda en cargar o aparece un error, espera unos segundos y "
        "recarga la página. Si persiste, contacta con Piensaenweb."
    ), tipo="NOTA")

    heading_2(doc, "Descripción rápida del panel")

    body(doc, (
        "Una vez dentro, verás un menú lateral con las secciones principales. "
        "Las que más usarás son:"
    ))

    bullet_list(doc, [
        "Productos (pim_products) — catálogo de platos de tiro.",
        "Blog / Noticias (blog_posts) — artículos publicados en la web.",
        "Contactos (crm_leads) — mensajes recibidos desde el formulario de contacto.",
    ])

    nota(doc, (
        "El resto de secciones son de configuración técnica. "
        "No las toques salvo que Piensaenweb te indique lo contrario."
    ), tipo="ADVERTENCIA")

    doc.add_page_break()

# ── Sección 2: Editar producto ─────────────────────────────────────────────

def add_seccion_productos(doc):
    heading_1(doc, "2. Cómo editar un producto existente")

    body(doc, (
        "Desde aquí puedes actualizar la información de cualquier plato de tiro: "
        "nombre, descripción, imágenes, especificaciones técnicas y más."
    ))

    heading_2(doc, "Abrir la lista de productos")

    step_list(doc, [
        "En el menú lateral, haz clic en «pim_products» (o «Productos» si lo ves traducido).",
        "Verás una tabla con todos los productos ordenados por nombre.",
        "Haz clic sobre el producto que quieres editar.",
    ])

    captura_placeholder(doc, "Lista de productos — tabla con nombre, estado y gama")

    heading_2(doc, "Campos que puedes editar")

    campo_tabla(doc, [
        ("name",              "Nombre del producto en español (campo principal)."),
        ("subtitle",          "Frase corta descriptiva que aparece bajo el nombre."),
        ("description_short", "Resumen breve para listados y tarjetas de producto."),
        ("description",       "Descripción completa que aparece en la página de detalle."),
        ("image",             "Fotografía principal del producto (fondo de color)."),
        ("image_transparent", "Fotografía del producto con fondo transparente (PNG). "
                              "Se usa en las tarjetas del catálogo."),
        ("status",            "Estado de publicación: «published» = visible en la web, "
                              "«draft» = borrador (no visible)."),
        ("featured",          "Marcar si el producto debe aparecer destacado en la portada."),
    ])

    nota(doc, (
        "Los campos de especificaciones técnicas (diameter_mm, weight_g, height_mm, etc.) "
        "son datos numéricos precisos. Modifícalos solo si tienes la información correcta "
        "del departamento técnico."
    ))

    heading_2(doc, "Cómo cambiar la imagen de un producto")

    step_list(doc, [
        "Dentro del producto, localiza el campo «image» o «image_transparent».",
        "Haz clic en la imagen actual (o en el botón «Seleccionar» si está vacío).",
        "En el explorador de archivos que se abre, puedes subir una imagen nueva "
        "(arrástrala o usa el botón «Subir archivo») o seleccionar una ya existente.",
        "Confirma la selección haciendo clic en «Confirmar».",
        "Verás la nueva imagen en el campo.",
        "Haz clic en «Guardar» (icono del disco o botón azul arriba a la derecha).",
    ])

    captura_placeholder(doc, "Formulario de edición de producto — vista general de campos")
    captura_placeholder(doc, "Explorador de archivos para subir o seleccionar imagen")

    nota(doc, (
        "Formatos de imagen recomendados: JPG para fotos con fondo, PNG para fotos "
        "con fondo transparente. Tamaño máximo recomendado: 2 MB. "
        "El sistema optimiza las imágenes automáticamente."
    ))

    heading_2(doc, "Guardar los cambios")

    step_list(doc, [
        "Una vez hayas terminado de editar, haz clic en el botón «Guardar» "
        "(parte superior derecha, icono de disco o «Save»).",
        "Verás un mensaje de confirmación breve en la parte inferior de la pantalla.",
        "Los cambios se reflejan en la web en pocos segundos.",
    ])

    nota(doc, (
        "Si cierras la pestaña sin guardar, perderás los cambios. "
        "Directus no guarda automáticamente."
    ), tipo="ADVERTENCIA")

    doc.add_page_break()

# ── Sección 3: Nueva noticia ───────────────────────────────────────────────

def add_seccion_noticias(doc):
    heading_1(doc, "3. Cómo publicar una noticia nueva")

    body(doc, (
        "El blog de Vivaz es una herramienta para comunicar novedades, "
        "participación en ferias, cambios de normativa y cualquier otro contenido "
        "de interés para clientes y distribuidores."
    ))

    heading_2(doc, "Crear una noticia nueva")

    step_list(doc, [
        "En el menú lateral, haz clic en «blog_posts» (o «Noticias / Blog»).",
        "Haz clic en el botón «+ Crear elemento» (parte superior derecha).",
        "Se abrirá un formulario en blanco.",
    ])

    captura_placeholder(doc, "Lista de noticias con el botón «+ Crear elemento» visible")

    heading_2(doc, "Campos de la noticia")

    campo_tabla(doc, [
        ("title",           "Título principal de la noticia (en español)."),
        ("slug",            "Identificador único para la URL. Ejemplo: «vivaz-en-feria-caza-2026». "
                            "Solo letras minúsculas, números y guiones. Sin espacios ni tildes."),
        ("excerpt",         "Resumen de 1-2 frases que aparece en el listado del blog."),
        ("content",         "Cuerpo completo de la noticia. Admite texto enriquecido "
                            "(negritas, listas, enlaces, etc.)."),
        ("image",           "Imagen destacada que aparece en la cabecera de la noticia "
                            "y en las tarjetas del listado."),
        ("category",        "Categoría de la noticia (por ejemplo: «Normativa», «Eventos»)."),
        ("published_at",    "Fecha de publicación. Puedes poner una fecha pasada o futura. "
                            "Si es futura, la noticia no será visible hasta esa fecha."),
        ("status",          "«published» = visible en la web. «draft» = borrador."),
        ("seo_title",       "Título para buscadores (Google). Si lo dejas vacío, "
                            "se usa el título principal."),
        ("seo_description", "Descripción breve para Google (máx. 160 caracteres)."),
    ])

    nota(doc, (
        "El slug es importante: define la URL de la noticia. Una vez publicada la noticia "
        "y compartida en redes sociales, evita cambiarlo para no romper enlaces existentes."
    ))

    heading_2(doc, "Guardar y publicar")

    step_list(doc, [
        "Rellena al menos los campos: título, slug, contenido e imagen.",
        "Cambia el campo «status» a «published».",
        "Pon la fecha en «published_at» (hoy o la fecha que corresponda).",
        "Haz clic en «Guardar».",
        "La noticia aparecerá en la web en la sección de noticias.",
    ])

    nota(doc, (
        "Si quieres preparar la noticia sin publicarla todavía, guárdala con status «draft». "
        "Puedes volver a editarla en cualquier momento y cambiar el estado a «published» "
        "cuando estés listo."
    ))

    captura_placeholder(doc, "Formulario de creación de noticia — campos principales")
    captura_placeholder(doc, "Editor de contenido enriquecido con opciones de formato")

    doc.add_page_break()

# ── Sección 4: Traducciones ────────────────────────────────────────────────

def add_seccion_traducciones(doc):
    heading_1(doc, "4. Cómo gestionar las traducciones")

    body(doc, (
        "La web de Vivaz está disponible en cuatro idiomas: español, inglés, francés y alemán. "
        "Cada contenido (producto o noticia) tiene su propio apartado de traducciones "
        "dentro del mismo formulario de edición."
    ))

    heading_2(doc, "Cómo llegar al apartado de traducciones")

    step_list(doc, [
        "Abre el producto o la noticia que quieres traducir (edita un elemento existente).",
        "Desplázate hacia abajo en el formulario.",
        "Encontrarás un bloque llamado «Translations» o «Traducciones».",
        "Dentro verás pestañas o entradas para cada idioma: «es» (español), «en» (inglés), "
        "«fr» (francés), «de» (alemán).",
    ])

    captura_placeholder(doc, "Bloque de traducciones dentro del formulario de un producto")

    heading_2(doc, "Cómo añadir o editar la traducción de un idioma")

    step_list(doc, [
        "Haz clic en el idioma que quieres editar (por ejemplo, «en» para inglés).",
        "Si ya existe traducción, verás los campos rellenos. Si no, estarán en blanco.",
        "Edita los campos de texto: nombre, subtítulo, descripción corta y descripción completa.",
        "Haz clic en «Guardar» para confirmar los cambios.",
    ])

    campo_tabla(doc, [
        ("name",              "Nombre del producto o título de la noticia en el idioma seleccionado."),
        ("subtitle",          "Subtítulo o eslogan (productos)."),
        ("description_short", "Resumen breve para listados."),
        ("description",       "Texto completo."),
        ("seo_title",         "Título para buscadores en ese idioma (noticias)."),
        ("seo_description",   "Descripción para buscadores en ese idioma (noticias)."),
    ])

    nota(doc, (
        "No es obligatorio tener todos los idiomas traducidos. Si un idioma no tiene traducción, "
        "la web mostrará el contenido en español como alternativa. "
        "Sin embargo, para el mercado de exportación es muy recomendable tener al menos "
        "inglés y alemán completos."
    ))

    heading_2(doc, "Recomendaciones para las traducciones")

    bullet_list(doc, [
        "Inglés (en): prioritario para exportación internacional.",
        "Alemán (de): prioritario para el mercado DACH (Alemania, Austria, Suiza).",
        "Francés (fr): importante para mercado francófono y algunos clubes europeos.",
        "Español (es): es el contenido principal — siempre debe estar completo.",
        "Para traducciones profesionales, puedes usar DeepL (deepl.com) como punto de partida "
        "y luego revisar el texto antes de pegarlo en Directus.",
    ])

    captura_placeholder(doc, "Formulario de traducción al inglés de un producto — campos rellenos")

    nota(doc, (
        "Si usas un traductor automático, revisa siempre la terminología técnica de tiro "
        "(trap, skeet, sporting, etc.) para asegurarte de que es correcta en el idioma destino."
    ), tipo="NOTA")

    doc.add_page_break()

# ── Sección 5: Leads / CRM ────────────────────────────────────────────────

def add_seccion_leads(doc):
    heading_1(doc, "5. Cómo ver los contactos recibidos")

    body(doc, (
        "Cada vez que alguien rellena el formulario de contacto de tu web, "
        "el mensaje queda registrado automáticamente en Directus. "
        "También recibirás un email de notificación, pero desde aquí puedes "
        "consultar el historial completo de contactos."
    ))

    heading_2(doc, "Acceder a los contactos")

    step_list(doc, [
        "En el menú lateral, haz clic en «crm_leads» (o «Contactos»).",
        "Verás una lista con todos los contactos recibidos, ordenados del más reciente al más antiguo.",
        "Haz clic sobre cualquier contacto para ver todos sus datos.",
    ])

    captura_placeholder(doc, "Lista de contactos (crm_leads) con nombre, email y fecha")

    heading_2(doc, "Información disponible de cada contacto")

    campo_tabla(doc, [
        ("name",      "Nombre de la persona que ha contactado."),
        ("apellidos", "Apellidos."),
        ("email",     "Correo electrónico de contacto."),
        ("phone",     "Teléfono de contacto."),
        ("company",   "Empresa u organización (si lo ha indicado)."),
        ("market",    "Mercado: «Nacional» o «Internacional»."),
        ("interest",  "Tipo de interés: «Distribución», «Club/Campo» o «Tirador»."),
        ("message",   "Mensaje completo que ha enviado."),
        ("source_page", "Página de la web desde la que envió el formulario."),
    ])

    nota(doc, (
        "Esta sección es de solo consulta. No es necesario que modifiques ni elimines "
        "ningún registro. Si quieres exportar los contactos a Excel, "
        "contacta con Piensaenweb y te preparamos el proceso."
    ))

    heading_2(doc, "Filtrar y buscar contactos")

    step_list(doc, [
        "En la parte superior de la lista, verás opciones de filtro y búsqueda.",
        "Puedes filtrar por «market» (Nacional / Internacional) o por «interest» "
        "(Distribución / Club / Tirador).",
        "También puedes usar la barra de búsqueda para encontrar un contacto por nombre o email.",
    ])

    captura_placeholder(doc, "Filtros de búsqueda en la lista de contactos")

    doc.add_page_break()

# ── Sección 6: Qué NO tocar ───────────────────────────────────────────────

def add_seccion_no_tocar(doc):
    heading_1(doc, "6. Qué NO debes tocar")

    body(doc, (
        "Directus contiene algunas colecciones y configuraciones técnicas que gestionan "
        "el funcionamiento interno de la web. Modificarlas sin conocimiento técnico "
        "puede causar errores en el sitio."
    ))

    heading_2(doc, "Colecciones que no debes modificar")

    campo_tabla(doc, [
        ("sys_brand",         "Configuración de marca: colores, tipografías, logo. "
                              "Solo Piensaenweb la modifica."),
        ("web_regulation",    "Datos de la normativa 2026 (fecha límite, límites PAH). "
                              "Requiere verificación legal antes de cualquier cambio."),
        ("web_videos",        "Gestión de vídeos embebidos. Contacta con Piensaenweb si "
                              "necesitas añadir o quitar vídeos."),
        ("pim_disciplines",   "Disciplinas de tiro (trap, skeet, sporting...). "
                              "No añadas ni elimines disciplinas sin consultarlo."),
        ("directus_users",    "Gestión de usuarios del panel. No crees ni borres usuarios."),
        ("directus_files",    "Biblioteca de archivos. Puedes subir imágenes, "
                              "pero no elimines archivos existentes sin confirmar que ya no se usan."),
    ])

    nota(doc, (
        "Si necesitas hacer algún cambio en estas secciones, escríbenos a "
        "claudio@piensaenweb.com y lo gestionamos juntos para evitar cualquier incidencia."
    ), tipo="ADVERTENCIA")

    heading_2(doc, "Configuraciones de sistema que no debes tocar")

    bullet_list(doc, [
        "Ajustes del proyecto (Settings → Project Settings).",
        "Roles y permisos (Settings → Roles & Permissions).",
        "Webhooks y flujos automatizados (Settings → Flows / Webhooks).",
        "Configuración de correo saliente (Settings → Email).",
    ])

    captura_placeholder(doc, "Menú de ajustes de sistema — las opciones marcadas en rojo no deben tocarse")

    nota(doc, (
        "Una regla sencilla: si no sabes exactamente para qué sirve una opción, "
        "no la cambies. En caso de duda, siempre es mejor preguntar primero."
    ))

    doc.add_page_break()

# ── Sección 7: FAQ y soporte ───────────────────────────────────────────────

def add_seccion_soporte(doc):
    heading_1(doc, "7. Preguntas frecuentes y contacto de soporte")

    heading_2(doc, "Preguntas frecuentes")

    preguntas = [
        (
            "He guardado los cambios pero no los veo en la web. ¿Por qué?",
            "La web puede tardar hasta 60 segundos en reflejar los cambios debido al sistema de caché. "
            "Si pasado ese tiempo no aparecen, recarga la página web con Ctrl+F5 (o Cmd+Shift+R en Mac). "
            "Si sigue sin aparecer, contacta con Piensaenweb."
        ),
        (
            "¿Puedo añadir un producto nuevo yo mismo?",
            "Sí, pero te recomendamos que los primeros productos nuevos los hagas junto a Piensaenweb "
            "para asegurarte de rellenar todos los campos técnicos correctamente "
            "(especificaciones, disciplinas, datos logísticos, etc.)."
        ),
        (
            "He subido una imagen y se ve cortada o pixelada. ¿Qué hago?",
            "Asegúrate de que la imagen tiene al menos 1200 píxeles de ancho y un peso inferior a 2 MB. "
            "Si la imagen es correcta pero sigue viéndose mal, contacta con Piensaenweb."
        ),
        (
            "¿Puedo borrar una noticia o un producto?",
            "Técnicamente sí, pero no lo recomendamos. Es mejor cambiar el estado a «draft» "
            "para que deje de ser visible sin perder los datos. "
            "Si necesitas eliminar algo definitivamente, consúltanos antes."
        ),
        (
            "He olvidado mi contraseña. ¿Cómo la recupero?",
            "En la pantalla de inicio de sesión, haz clic en «¿Olvidaste tu contraseña?» "
            "e introduce tu email. Recibirás un enlace para restablecerla. "
            "Si no llega el email, contacta con Piensaenweb."
        ),
        (
            "¿Puedo acceder desde el móvil?",
            "Sí, Directus tiene una interfaz adaptable al móvil. Para ediciones largas "
            "es más cómodo usar un ordenador, pero para consultas rápidas o cambios sencillos "
            "el móvil funciona bien."
        ),
    ]

    for pregunta, respuesta in preguntas:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(f"P: {pregunta}")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_VERDE

        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after  = Pt(6)
        p2.paragraph_format.left_indent  = Cm(0.5)
        run2 = p2.add_run(f"R: {respuesta}")
        run2.font.size = Pt(11)
        run2.font.color.rgb = COLOR_GRIS_TEXTO

    heading_2(doc, "Contacto de soporte técnico")

    body(doc, "Para cualquier duda o incidencia relacionada con la web o el panel de Directus:")

    campo_tabla(doc, [
        ("Agencia",   "Piensaenweb"),
        ("Email",     "claudio@piensaenweb.com"),
        ("Web",       "piensaenweb.com"),
        ("Horario",   "Lunes a viernes, 9:00 – 18:00 h (hora española)"),
    ])

    nota(doc, (
        "Cuando nos escribas con una incidencia, incluye: una descripción de lo que querías hacer, "
        "qué ha pasado exactamente y, si puedes, una captura de pantalla del error. "
        "Así te ayudamos mucho más rápido."
    ))

    # Pie de pagina
    doc.add_paragraph()
    add_horizontal_line(doc, "CCCCCC")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"Documento preparado por Piensaenweb · Mayo 2026 · "
        f"Para uso exclusivo de Vivaz Clay Targets"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run.italic = True

# ── Main ───────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    # Margenes del documento
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Fuente por defecto
    style = doc.styles['Normal']
    font  = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Contenido
    add_portada(doc)
    add_indice(doc)
    add_seccion_acceso(doc)
    add_seccion_productos(doc)
    add_seccion_noticias(doc)
    add_seccion_traducciones(doc)
    add_seccion_leads(doc)
    add_seccion_no_tocar(doc)
    add_seccion_soporte(doc)

    doc.save(OUTPUT_PATH)
    print(f"Documento guardado en:\n{OUTPUT_PATH}")

if __name__ == "__main__":
    main()
